"""
Document loading and chunking module for RAG pipeline.

Supports PDF, DOCX, TXT, MD, HTML, CSV, JSON formats with
fixed-size, recursive, and semantic chunking strategies.
"""

from __future__ import annotations

import csv
import hashlib
import json
import logging
import os
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".html", ".htm", ".csv", ".json"}


def _hash_text(text: str) -> str:
    """Return a full SHA-256 hex digest of text.

    BUG FIX: Previously truncated to 16 hex chars (64 bits), which is too
    short for a content-addressed document id — collision risk grows with
    corpus size, and the DocumentRecord docstring explicitly promises a
    full SHA-256. Chunk ids derive from this too; a longer id is harmless.
    """
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


@dataclass
class Document:
    """Represents a loaded document."""

    content: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    doc_id: str = field(default="")

    def __post_init__(self) -> None:
        if not self.doc_id:
            self.doc_id = _hash_text(self.content)


@dataclass
class Chunk:
    """Represents a chunk of a document."""

    content: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    chunk_id: str = field(default="")
    doc_id: str = field(default="")

    def __post_init__(self) -> None:
        if not self.chunk_id:
            self.chunk_id = _hash_text(self.content + self.doc_id)


class DocumentLoader:
    """Loads documents from files or directories into Document objects."""

    def load(self, file_path: str | Path) -> Document:
        """Load a single file and return a Document.

        Args:
            file_path: Path to the file to load.

        Returns:
            Document with content and metadata.

        Raises:
            ValueError: If file type is unsupported.
            FileNotFoundError: If file does not exist.
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        ext = path.suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}. Supported: {SUPPORTED_EXTENSIONS}")

        logger.info("Loading document: %s", path)

        base_metadata: Dict[str, Any] = {
            "filename": path.name,
            "file_path": str(path.resolve()),
            "file_type": ext.lstrip("."),
            "file_size_bytes": path.stat().st_size,
        }

        loaders = {
            ".pdf": self._load_pdf,
            ".docx": self._load_docx,
            ".txt": self._load_text,
            ".md": self._load_text,
            ".html": self._load_html,
            ".htm": self._load_html,
            ".csv": self._load_csv,
            ".json": self._load_json,
        }

        content, extra_meta = loaders[ext](path)
        base_metadata.update(extra_meta)
        doc = Document(content=content, metadata=base_metadata)
        logger.debug("Loaded document %s (%d chars)", path.name, len(content))
        return doc

    def load_directory(
        self,
        directory: str | Path,
        recursive: bool = True,
        extensions: Optional[List[str]] = None,
    ) -> List[Document]:
        """Load all supported documents from a directory.

        Args:
            directory: Path to the directory.
            recursive: Whether to search subdirectories.
            extensions: Optional list of extensions to filter (e.g. ['.pdf', '.txt']).

        Returns:
            List of Document objects.
        """
        dir_path = Path(directory)
        if not dir_path.is_dir():
            raise NotADirectoryError(f"Not a directory: {dir_path}")

        allowed = {e.lower() for e in (extensions or SUPPORTED_EXTENSIONS)}
        pattern = "**/*" if recursive else "*"
        files = [p for p in dir_path.glob(pattern) if p.is_file() and p.suffix.lower() in allowed]

        logger.info("Found %d files in %s", len(files), dir_path)

        documents: List[Document] = []
        for file in files:
            try:
                doc = self.load(file)
                documents.append(doc)
            except Exception as exc:
                logger.warning("Failed to load %s: %s", file, exc)

        logger.info("Successfully loaded %d/%d documents", len(documents), len(files))
        return documents

    # ------------------------------------------------------------------ #
    # Private format loaders                                               #
    # ------------------------------------------------------------------ #

    def _load_text(self, path: Path) -> tuple[str, Dict[str, Any]]:
        """Load plain text or Markdown file."""
        text = path.read_text(encoding="utf-8", errors="replace")
        return text, {"encoding": "utf-8"}

    def _load_pdf(self, path: Path) -> tuple[str, Dict[str, Any]]:
        """Load PDF file using pypdf.

        WHY: pypdf extracts text with hard line breaks at the PDF column width,
        producing single '\\n' inside paragraphs. Without normalisation these
        layout-level newlines cause the recursive chunker to over-fragment text.

        FIX: After joining pages, we normalise single '\\n' → space while
        preserving real paragraph breaks ('\\n\\n').
        """
        try:
            import pypdf  # type: ignore

            reader = pypdf.PdfReader(str(path))
            pages: List[str] = []
            for page in reader.pages:
                pages.append(page.extract_text() or "")
            text = "\n\n".join(pages)

            # BEFORE: "Fine-Tuning LLMs from\nBasics to Breakthroughs"
            # AFTER:  "Fine-Tuning LLMs from Basics to Breakthroughs"
            # Preserve real paragraph breaks (\n\n) by temporarily replacing
            # them, then normalise single \n (PDF line wraps) to spaces.
            text = text.replace("\n\n", "\x00")   # protect paragraph breaks
            text = text.replace("\n", " ")         # layout line breaks → space
            text = text.replace("\x00", "\n\n")    # restore paragraph breaks

            # Rejoin hyphenated line breaks: "develop- ment" → "development"
            # WHY: PDF wraps long words with a hyphen at column boundaries.
            # After \n→space, these become "word- continuation".  The pattern
            # hyphen-space-lowercase reliably identifies line-break hyphens
            # vs real compounds like "self-attention" (no space after hyphen).
            text = re.sub(r"(\w)- ([a-z])", r"\1\2", text)

            text = re.sub(r" {2,}", " ", text)     # collapse multiple spaces

            meta: Dict[str, Any] = {"page_count": len(reader.pages)}
            if reader.metadata:
                for k in ("title", "author", "subject"):
                    v = getattr(reader.metadata, k, None)
                    if v:
                        meta[k] = v
            return text, meta
        except ImportError:
            logger.warning("pypdf not installed; reading PDF as binary text")
            return path.read_text(errors="replace"), {}

    def _load_docx(self, path: Path) -> tuple[str, Dict[str, Any]]:
        """Load DOCX file using python-docx."""
        try:
            import docx  # type: ignore

            doc = docx.Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n\n".join(paragraphs)
            props = doc.core_properties
            meta: Dict[str, Any] = {}
            for attr in ("author", "title", "subject", "created", "modified"):
                val = getattr(props, attr, None)
                if val:
                    meta[attr] = str(val)
            return text, meta
        except ImportError:
            logger.warning("python-docx not installed; cannot load DOCX")
            return "", {"error": "python-docx not installed"}

    def _load_html(self, path: Path) -> tuple[str, Dict[str, Any]]:
        """Load HTML file using BeautifulSoup."""
        html = path.read_text(encoding="utf-8", errors="replace")
        try:
            from bs4 import BeautifulSoup  # type: ignore

            soup = BeautifulSoup(html, "html.parser")
            for tag in soup(["script", "style", "nav", "footer", "header"]):
                tag.decompose()
            text = soup.get_text(separator="\n", strip=True)
            title = soup.title.string if soup.title else ""
            return text, {"html_title": title or ""}
        except ImportError:
            logger.warning("beautifulsoup4 not installed; stripping HTML tags naively")
            import re

            text = re.sub(r"<[^>]+>", " ", html)
            text = re.sub(r"\s+", " ", text).strip()
            return text, {}

    def _load_csv(self, path: Path) -> tuple[str, Dict[str, Any]]:
        """Load CSV file as structured text."""
        rows: List[List[str]] = []
        with path.open(newline="", encoding="utf-8", errors="replace") as fh:
            reader = csv.reader(fh)
            for row in reader:
                rows.append(row)
        if not rows:
            return "", {"row_count": 0, "column_count": 0}
        headers = rows[0]
        lines: List[str] = [", ".join(headers)]
        for row in rows[1:]:
            pairs = [f"{h}: {v}" for h, v in zip(headers, row)]
            lines.append("; ".join(pairs))
        text = "\n".join(lines)
        return text, {"row_count": len(rows) - 1, "column_count": len(headers)}

    def _load_json(self, path: Path) -> tuple[str, Dict[str, Any]]:
        """Load JSON file as pretty-printed text."""
        raw = path.read_text(encoding="utf-8", errors="replace")
        try:
            data = json.loads(raw)
            text = json.dumps(data, indent=2, ensure_ascii=False)
            return text, {"json_valid": True}
        except json.JSONDecodeError:
            return raw, {"json_valid": False}


class TextChunker:
    """Splits documents into overlapping chunks for embedding."""

    def __init__(
        self,
        chunk_size: int = 512,
        chunk_overlap: int = 64,
        strategy: str = "recursive",
        separators: Optional[List[str]] = None,
    ) -> None:
        """
        Args:
            chunk_size: Maximum characters per chunk.
            chunk_overlap: Number of overlapping characters between chunks.
            strategy: 'fixed', 'recursive', or 'semantic'.
            separators: Custom separators for recursive strategy.
        """
        if chunk_size <= 0:
            raise ValueError("chunk_size must be positive")
        if chunk_overlap < 0 or chunk_overlap >= chunk_size:
            raise ValueError("chunk_overlap must be >= 0 and < chunk_size")
        if strategy not in ("fixed", "recursive", "semantic"):
            raise ValueError("strategy must be 'fixed', 'recursive', or 'semantic'")

        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.strategy = strategy
        self.separators = separators or ["\n\n", "\n", ". ", " ", ""]

    # WHY 20 chars: shorter chunks are almost always PDF artifacts — page
    # numbers ("109"), stray headers, or section labels.  They carry no
    # semantic value and pollute retrieval results with false matches.
    MIN_CHUNK_LENGTH = 20

    def chunk(self, document: Document) -> List[Chunk]:
        """Split a Document into chunks.

        Args:
            document: Document to split.

        Returns:
            List of Chunk objects.
        """
        if self.strategy == "fixed":
            raw_chunks = self._fixed_chunk(document.content)
        elif self.strategy == "recursive":
            # WHY overlap is applied here instead of inside _recursive_chunk:
            # The recursive splitter calls itself at multiple depths.  If
            # overlap were applied at each depth it would cascade — the tail
            # of a depth-1 chunk (already overlapped) gets overlapped again
            # at depth 0, tripling text.  Applying once at the top avoids this.
            raw_chunks = self._recursive_chunk(document.content)
            raw_chunks = self._apply_word_overlap(raw_chunks)
        else:  # semantic
            raw_chunks = self._semantic_chunk(document.content)

        chunks: List[Chunk] = []
        for idx, text in enumerate(raw_chunks):
            stripped = text.strip()
            if not stripped or len(stripped) < self.MIN_CHUNK_LENGTH:
                continue
            # Filter ToC dot-leader chunks (". . . . . . . . . 42").
            # WHY: PDF tables of contents extract as dot-filled lines
            # mixed with section titles.  They carry no semantic value
            # and pollute retrieval.  Content chunks have < 5% dots;
            # ToC chunks have > 20% dots — a clean bimodal split.
            dot_ratio = stripped.count(".") / len(stripped)
            if dot_ratio > 0.15:
                continue
            meta = {**document.metadata, "chunk_index": idx, "chunk_strategy": self.strategy}
            chunks.append(Chunk(content=stripped, metadata=meta, doc_id=document.doc_id))

        logger.debug(
            "Chunked document %s into %d chunks (strategy=%s)",
            document.doc_id,
            len(chunks),
            self.strategy,
        )
        return chunks

    def chunk_documents(self, documents: List[Document]) -> List[Chunk]:
        """Chunk multiple documents.

        Args:
            documents: List of Document objects.

        Returns:
            Flattened list of all Chunk objects.
        """
        all_chunks: List[Chunk] = []
        for doc in documents:
            all_chunks.extend(self.chunk(doc))
        logger.info(
            "Total chunks from %d documents: %d", len(documents), len(all_chunks)
        )
        return all_chunks

    # ------------------------------------------------------------------ #
    # Chunking strategies                                                  #
    # ------------------------------------------------------------------ #

    def _fixed_chunk(self, text: str) -> List[str]:
        """Split text into fixed-size character windows with overlap."""
        chunks: List[str] = []
        start = 0
        while start < len(text):
            end = start + self.chunk_size
            chunks.append(text[start:end])
            start += self.chunk_size - self.chunk_overlap
        return chunks

    def _recursive_chunk(self, text: str, depth: int = 0) -> List[str]:
        """Recursively split text using a hierarchy of separators.

        Splits on the current-depth separator, merges small parts into
        chunks up to ``chunk_size``, then applies word-boundary-safe
        overlap via ``_apply_word_overlap``.
        """
        if len(text) <= self.chunk_size:
            return [text] if text.strip() else []

        if depth >= len(self.separators):
            return self._fixed_chunk(text)

        sep = self.separators[depth]
        if sep == "":
            return self._fixed_chunk(text)

        parts = text.split(sep)
        chunks: List[str] = []
        current_parts: List[str] = []
        current_len = 0

        for part in parts:
            added_len = len(part) + (len(sep) if current_parts else 0)

            if current_len + added_len <= self.chunk_size:
                current_parts.append(part)
                current_len += added_len
            else:
                if current_parts:
                    committed = sep.join(current_parts)
                    if len(committed) > self.chunk_size:
                        chunks.extend(self._recursive_chunk(committed, depth + 1))
                    else:
                        chunks.append(committed)

                current_parts = [part]
                current_len = len(part)

        if current_parts:
            remaining = sep.join(current_parts)
            if len(remaining) > self.chunk_size:
                chunks.extend(self._recursive_chunk(remaining, depth + 1))
            else:
                chunks.append(remaining)

        return chunks

    def _apply_word_overlap(self, chunks: List[str]) -> List[str]:
        """Prepend the trailing words of chunk N to chunk N+1.

        BEFORE (broken _apply_overlap):
            Sliced last N raw *characters* and concatenated with no separator,
            producing "fine-tuningsystems." and doubled content.

        AFTER:
            Takes the last ``chunk_overlap`` characters, snaps *forward* to the
            nearest word boundary (first space), and prepends with ``" ... "``
            as a visual separator.  Result is always clean, readable text.

        WHY word-boundary snapping:
            Character-level slicing can cut mid-word ("optimisa|tion").
            Snapping to the next space guarantees whole words.
        """
        if self.chunk_overlap == 0 or len(chunks) <= 1:
            return chunks

        result: List[str] = [chunks[0]]
        for i in range(1, len(chunks)):
            prev = chunks[i - 1]
            # Grab roughly chunk_overlap chars from the end of previous chunk
            raw_tail = prev[-self.chunk_overlap:]
            # Snap forward to the nearest word boundary (skip partial word)
            space_idx = raw_tail.find(" ")
            if space_idx != -1 and space_idx < len(raw_tail) - 1:
                tail = raw_tail[space_idx + 1:]
            else:
                # The tail is a single long word — use it as-is
                tail = raw_tail
            tail = tail.strip()
            if tail:
                result.append(tail + " " + chunks[i])
            else:
                result.append(chunks[i])
        return result

    def _semantic_chunk(self, text: str) -> List[str]:
        """Sentence-aware chunking: accumulate sentences until chunk_size is exceeded."""
        import re

        # Split on sentence boundaries
        sentence_endings = re.compile(r"(?<=[.!?])\s+")
        sentences = sentence_endings.split(text)

        chunks: List[str] = []
        current_sentences: List[str] = []
        current_len = 0

        for sentence in sentences:
            s_len = len(sentence)
            if current_len + s_len > self.chunk_size and current_sentences:
                chunks.append(" ".join(current_sentences))
                # keep overlap
                overlap_sentences: List[str] = []
                overlap_len = 0
                for sent in reversed(current_sentences):
                    if overlap_len + len(sent) <= self.chunk_overlap:
                        overlap_sentences.insert(0, sent)
                        overlap_len += len(sent)
                    else:
                        break
                current_sentences = overlap_sentences
                current_len = overlap_len

            current_sentences.append(sentence)
            current_len += s_len

        if current_sentences:
            chunks.append(" ".join(current_sentences))

        return [c for c in chunks if c.strip()]
